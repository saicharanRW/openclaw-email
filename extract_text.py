import sys
import argparse
from docx import Document
from docx.oxml import parse_xml
from krutidev_converter import krutidev_to_unicode

def extract_from_table(table, convert_krutidev):
    """Recursively extract text from a table and its cells (handling nested tables)."""
    table_text = []
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            # A cell can contain paragraphs and nested tables
            cell_parts = []
            for element in cell._element:
                from docx.table import Table
                from docx.text.paragraph import Paragraph
                
                if element.tag.endswith('p'):
                    para = Paragraph(element, cell)
                    text = para.text.strip()
                    if text:
                        if convert_krutidev:
                            text = krutidev_to_unicode(text)
                        cell_parts.append(text)
                elif element.tag.endswith('tbl'):
                    nested_table = Table(element, cell)
                    nested_text = extract_from_table(nested_table, convert_krutidev)
                    if nested_text:
                        cell_parts.append(f"[\n{nested_text}\n]")
            
            row_text.append('\n'.join(cell_parts))
        
        if any(cell_content.strip() for cell_content in row_text):
            table_text.append(' | '.join(row_text))
    
    return '\n'.join(table_text)

def extract_text(docx_path, convert_krutidev=False):
    """
    Extract ALL text from DOCX file including paragraphs, tables, headers, footer, 
    footnotes, and endnotes, preserving document order.
    """
    try:
        doc = Document(docx_path)
        full_text = []
        
        # Iterate through elements in the body to preserve order
        for element in doc.element.body:
            from docx.table import Table
            from docx.text.paragraph import Paragraph
            
            if element.tag.endswith('p'):
                para = Paragraph(element, doc)
                text = para.text.strip()
                if text:
                    if convert_krutidev:
                        text = krutidev_to_unicode(text)
                    full_text.append(text)
                    
            elif element.tag.endswith('tbl'):
                table = Table(element, doc)
                table_content = extract_from_table(table, convert_krutidev)
                if table_content:
                    full_text.append(table_content)
        
        # Extract from headers
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                text = para.text.strip()
                if text:
                    if convert_krutidev:
                        text = krutidev_to_unicode(text)
                    full_text.append(text)
        
        # Extract from footers
        for section in doc.sections:
            footer = section.footer
            for para in footer.paragraphs:
                text = para.text.strip()
                if text:
                    if convert_krutidev:
                        text = krutidev_to_unicode(text)
                    full_text.append(text)
        
        # Extract from footnotes if they exist
        try:
            if hasattr(doc.part, 'footnotes') and doc.part.footnotes is not None:
                for footnote in doc.part.footnotes.footnotes:
                    for para in footnote.paragraphs:
                        text = para.text.strip()
                        if text:
                            if convert_krutidev:
                                text = krutidev_to_unicode(text)
                            full_text.append(f"[Footnote: {text}]")
        except:
            pass  # Skip if footnotes are not available
        
        # Extract from endnotes if they exist
        try:
            if hasattr(doc.part, 'endnotes') and doc.part.endnotes is not None:
                for endnote in doc.part.endnotes.endnotes:
                    for para in endnote.paragraphs:
                        text = para.text.strip()
                        if text:
                            if convert_krutidev:
                                text = krutidev_to_unicode(text)
                            full_text.append(f"[Endnote: {text}]")
        except:
            pass  # Skip if endnotes are not available
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error reading {docx_path}: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description='Extract text from DOCX files, optionally converting Krutidev to Unicode',
        prog='extract_text'
    )
    parser.add_argument('docx_file', help='Input DOCX file')
    parser.add_argument('output_file', nargs='?', help='Output file (optional, prints to stdout if not provided)')
    parser.add_argument('--convert', action='store_true', help='Convert Krutidev text to Unicode')
    
    args = parser.parse_args()
    
    text = extract_text(args.docx_file, convert_krutidev=args.convert)
    
    if args.output_file:
        try:
            with open(args.output_file, 'w', encoding='utf-8') as f:
                f.write(text)
            print(f"Extracted text written to {args.output_file}")
        except Exception as e:
            print(f"Error writing to {args.output_file}: {e}", file=sys.stderr)
            sys.exit(1)
    else:
        # Print to stdout, using utf-8 explicitly to avoid encoding issues
        try:
            sys.stdout.reconfigure(encoding='utf-8')
        except:
            pass
        print(text)
