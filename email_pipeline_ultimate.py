"""
email_pipeline_enhanced.py
---------------------------
Enhanced unified pipeline that:
  1. Polls inbox for new emails with attachments (PDF, DOCX, XLSX, CSV)
  2. Saves attachments locally
  3. Extracts text using document indexer's extraction functions
  4. Sends extracted text to OpenAI to find tables and extract column headers
  5. Sends email context + column headers to OpenClaw to get structured data
  6. Creates reply attachment in THE SAME FORMAT as received (PDF→PDF, DOCX→DOCX, etc.)
  7. Sends auto-reply with the generated attachment

Supports all document types from document_indexer_ULTIMATE.py:
  - PDF
  - DOCX (with Krutidev support)
  - XLSX
  - CSV

Requires:
    pip install pyzmail36 requests openai reportlab pypdf \
                python-docx openpyxl PyPDF2
    
All configuration lives in config.py.

FONT SETUP (required for Hindi/Devanagari PDF output):
  Place NotoSansDevanagari-Regular.ttf next to this file, OR install it in
  C:\\Windows\\Fonts\\ on Windows / /usr/share/fonts/ on Linux.
  Download: https://fonts.google.com/noto/specimen/Noto+Sans+Devanagari
"""

import datetime
import imaplib
import json
import os
import re
import smtplib
import time
import uuid
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

from openai import OpenAI
import pyzmail
import requests
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Import extraction functions from document indexer
import csv as csv_module
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
import PyPDF2

from config import (
    # IMAP
    IMAP_HOST, EMAIL, PASSWORD,
    # SMTP
    SMTP_SERVER, SMTP_PORT, SMTP_EMAIL, SMTP_PASSWORD,
    # OpenClaw
    OPENCLAW_URL, TOKEN,
    # OpenAI
    OPENAI_API_KEY, OPENAI_MODEL,
    # Misc
    CHECK_INTERVAL,
    ATTACHMENT_DIR,
    PROCESSED_LOG,
)

# Try to import Krutidev converter
try:
    from krutidev_converter import krutidev_to_unicode
    KRUTIDEV_AVAILABLE = True
except:
    KRUTIDEV_AVAILABLE = False
    krutidev_to_unicode = lambda x: x

# Try to import custom DOCX extractor if available
try:
    from extract_docx_text import extract_text as extract_docx_krutidev
    CUSTOM_DOCX_EXTRACTOR = True
except:
    CUSTOM_DOCX_EXTRACTOR = False
    extract_docx_krutidev = None

# ══════════════════════════════════════════════════════════════════════════════
# SUPPORTED FILE TYPES
# ══════════════════════════════════════════════════════════════════════════════

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".csv"}

# ══════════════════════════════════════════════════════════════════════════════
# DATABASE PATH  (single config point — passed into every OpenClaw prompt)
# ══════════════════════════════════════════════════════════════════════════════

DB_PATH = os.environ.get(
    "DOCUMENT_INDEX_DB",
    "/home/randomwalk/.openclaw/workspace/multilingual-document-processor/scripts/document_index.db",
)

# ══════════════════════════════════════════════════════════════════════════════
# OPENAI SETUP
# ══════════════════════════════════════════════════════════════════════════════

openai_client = OpenAI(api_key=OPENAI_API_KEY)

# ══════════════════════════════════════════════════════════════════════════════
# LOGGING
# ══════════════════════════════════════════════════════════════════════════════

def log(msg: str) -> None:
    ts = datetime.datetime.now().strftime("%H:%M:%S")
    print(f"[{ts}] {msg}")


# ══════════════════════════════════════════════════════════════════════════════
# PROCESSED FILE TRACKING
# ══════════════════════════════════════════════════════════════════════════════

def load_processed() -> set:
    if not os.path.exists(PROCESSED_LOG):
        return set()
    with open(PROCESSED_LOG, "r", encoding="utf-8") as f:
        return set(line.strip() for line in f if line.strip())


def mark_processed(filename: str) -> None:
    with open(PROCESSED_LOG, "a", encoding="utf-8") as f:
        f.write(filename + "\n")


# ══════════════════════════════════════════════════════════════════════════════
# PDF FONT HANDLING  ← NEW: fixes Hindi black-box rendering in reply PDFs
# ══════════════════════════════════════════════════════════════════════════════

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

_NOTO_REGULAR_PATHS = [
    os.path.join(_SCRIPT_DIR, "NotoSansDevanagari-Regular.ttf"),
    os.path.join(_SCRIPT_DIR, "NotoSansDevanagari-VariableFont_wdth,wght.ttf"),
    os.path.join(_SCRIPT_DIR, "NotoSansDevanagari.ttf"),
    r"C:\Windows\Fonts\NotoSansDevanagari-Regular.ttf",
    r"C:\Windows\Fonts\NotoSansDevanagari.ttf",
    "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf",
    "/usr/share/fonts/noto/NotoSansDevanagari-Regular.ttf",
]
_NOTO_BOLD_PATHS = [
    os.path.join(_SCRIPT_DIR, "NotoSansDevanagari-Bold.ttf"),
    r"C:\Windows\Fonts\NotoSansDevanagari-Bold.ttf",
    "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Bold.ttf",
    os.path.join(_SCRIPT_DIR, "NotoSansDevanagari-Regular.ttf"),
    os.path.join(_SCRIPT_DIR, "NotoSansDevanagari-VariableFont_wdth,wght.ttf"),
]
_ARIALUNI_PATHS = [
    r"C:\Windows\Fonts\arialuni.ttf",
    os.path.join(_SCRIPT_DIR, "arialuni.ttf"),
    "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
]

_hindi_font_regular: str  = "Helvetica"
_hindi_font_bold:    str  = "Helvetica-Bold"
_hindi_fonts_loaded: bool = False


def _load_hindi_fonts() -> None:
    """Register a Devanagari-capable font with ReportLab (runs once at startup)."""
    global _hindi_font_regular, _hindi_font_bold, _hindi_fonts_loaded
    if _hindi_fonts_loaded:
        return

    def _first(paths):
        return next((p for p in paths if os.path.exists(p)), None)

    reg_path  = _first(_NOTO_REGULAR_PATHS)
    bold_path = _first(_NOTO_BOLD_PATHS)

    if reg_path:
        try:
            pdfmetrics.registerFont(TTFont("NotoSansDevanagari", reg_path))
            pdfmetrics.registerFont(TTFont("NotoSansDevanagari-Bold", bold_path or reg_path))
            _hindi_font_regular = "NotoSansDevanagari"
            _hindi_font_bold    = "NotoSansDevanagari-Bold"
            log(f"  ✓ Hindi font loaded: {reg_path}")
            _hindi_fonts_loaded = True
            return
        except Exception as e:
            log(f"  ⚠ NotoSansDevanagari load failed: {e}")

    uni_path = _first(_ARIALUNI_PATHS)
    if uni_path:
        try:
            pdfmetrics.registerFont(TTFont("ArialUnicode", uni_path))
            _hindi_font_regular = "ArialUnicode"
            _hindi_font_bold    = "ArialUnicode"
            log(f"  ✓ Hindi font fallback loaded: {uni_path}")
            _hindi_fonts_loaded = True
            return
        except Exception as e:
            log(f"  ⚠ ArialUnicode load failed: {e}")

    log(
        "  ✗ WARNING: No Devanagari font found — Hindi text will show as black boxes.\n"
        "    Fix: place NotoSansDevanagari-Regular.ttf next to this script.\n"
        "    Download: https://fonts.google.com/noto/specimen/Noto+Sans+Devanagari"
    )
    _hindi_fonts_loaded = True


def _contains_devanagari(text: str) -> bool:
    return any(0x0900 <= ord(c) <= 0x097F for c in text)


def _smart_paragraph(text: str, latin_style: ParagraphStyle,
                      hindi_style: ParagraphStyle) -> Paragraph:
    """Return a Paragraph with the Hindi font if the text contains Devanagari."""
    s = str(text)
    return Paragraph(s, hindi_style if _contains_devanagari(s) else latin_style)


# ══════════════════════════════════════════════════════════════════════════════
# ENCODING DETECTION (from document indexer)
# ══════════════════════════════════════════════════════════════════════════════

def detect_encoding(text: str) -> dict:
    """Intelligently detect the encoding of text."""
    if not text or len(text) < 5:
        return {'type': 'unknown', 'confidence': 0.0, 'needs_conversion': False, 'script': 'unknown'}
    
    sample = text[:1000]
    
    devanagari_count = sum(1 for c in sample if 0x0900 <= ord(c) <= 0x097F)
    latin_count = sum(1 for c in sample if (ord(c) >= 65 and ord(c) <= 90) or (ord(c) >= 97 and ord(c) <= 122))
    extended_ascii_count = sum(1 for c in sample if 128 <= ord(c) < 256)
    
    total_chars = len(sample.replace(' ', '').replace('\n', '').replace('\t', ''))
    if total_chars == 0:
        return {'type': 'empty', 'confidence': 1.0, 'needs_conversion': False, 'script': 'none'}
    
    devanagari_ratio = devanagari_count / total_chars
    latin_ratio = latin_count / total_chars
    extended_ascii_ratio = extended_ascii_count / total_chars
    
    krutidev_patterns = ['k', 'Dk', '[k', 'Xk', '?k', 'Pk', 'Nk', 'Tk', 'vk', 'bZ', 'ks', 'kS']
    krutidev_pattern_count = sum(sample.count(p) for p in krutidev_patterns)
    krutidev_score = krutidev_pattern_count / (total_chars / 10)
    
    if devanagari_ratio > 0.3:
        return {'type': 'unicode', 'confidence': 0.9, 'needs_conversion': False, 'script': 'devanagari'}
    elif krutidev_score > 0.5 or extended_ascii_ratio > 0.1:
        confidence = min(0.9, krutidev_score)
        return {'type': 'krutidev', 'confidence': confidence, 'needs_conversion': True, 'script': 'devanagari'}
    elif latin_ratio > 0.5:
        return {'type': 'english', 'confidence': 0.9, 'needs_conversion': False, 'script': 'latin'}
    elif latin_ratio > 0.2 and devanagari_ratio > 0.05:
        return {'type': 'mixed', 'confidence': 0.7, 'needs_conversion': False, 'script': 'mixed'}
    else:
        return {'type': 'unknown', 'confidence': 0.3, 'needs_conversion': False, 'script': 'unknown'}


def smart_convert(text: str) -> tuple[str, dict]:
    """Intelligently convert text if needed."""
    detection = detect_encoding(text)
    
    if not detection['needs_conversion']:
        log(f"    → {detection['type'].capitalize()} detected, no conversion needed")
        return text, detection
    
    if not KRUTIDEV_AVAILABLE:
        log(f"    ⚠ Krutidev detected but converter not available")
        return text, detection
    
    try:
        log(f"    ✓ Converting Krutidev to Unicode (confidence: {detection['confidence']:.2f})")
        converted = krutidev_to_unicode(text)
        if len(converted) < len(text) * 0.3:
            log(f"    ⚠ Conversion produced suspiciously short text, keeping original")
            return text, detection
        return converted, detection
    except Exception as e:
        log(f"    ✗ Conversion failed: {e}")
        return text, detection


# ══════════════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION (from document indexer)
# ══════════════════════════════════════════════════════════════════════════════

def extract_text_from_pdf(filepath: str) -> str | None:
    """Extract text from PDF with encoding detection."""
    try:
        text = ""
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if not text.strip():
            return ""
        
        converted, _ = smart_convert(text)
        return converted
        
    except Exception as e:
        log(f"    ERROR extracting PDF: {e}")
        return None


def extract_text_from_docx(filepath: str) -> str | None:
    """Extract text from DOCX with Krutidev support."""
    try:
        doc = docx.Document(filepath)
        all_text = "\n".join([p.text for p in doc.paragraphs])
        detection = detect_encoding(all_text)
        
        if detection['needs_conversion'] and CUSTOM_DOCX_EXTRACTOR:
            log(f"    Converting entire DOCX as Krutidev")
            all_text = extract_docx_krutidev(filepath)
            return all_text
        
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        
        return "\n".join(parts)
        
    except Exception as e:
        log(f"    ERROR extracting DOCX: {e}")
        return None


def extract_text_from_xlsx(filepath: str) -> str | None:
    """Extract text from XLSX with encoding detection."""
    try:
        wb = openpyxl.load_workbook(filepath, data_only=True)
        parts = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            parts.append(f"--- Sheet: {sheet_name} ---")
            
            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        val = str(cell.value)
                        row_values.append(val)
                    else:
                        row_values.append("")
                parts.append("\t".join(row_values))
        
        full_text = "\n".join(parts)
        converted, _ = smart_convert(full_text)
        return converted
        
    except Exception as e:
        log(f"    ERROR extracting XLSX: {e}")
        return None


def extract_text_from_csv(filepath: str) -> str | None:
    """Extract text from CSV with encoding detection."""
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv_module.reader(f)
            rows = [row for row in reader]
        
        full_text = "\n".join(", ".join(row) for row in rows)
        converted, _ = smart_convert(full_text)
        return converted
        
    except Exception as e:
        log(f"    ERROR extracting CSV: {e}")
        return None


def extract_document_content(filepath: str) -> tuple[str | None, str]:
    """Extract text from any supported document type."""
    extension = os.path.splitext(filepath)[1].lower()
    
    dispatch = {
        ".pdf":  extract_text_from_pdf,
        ".docx": extract_text_from_docx,
        ".xlsx": extract_text_from_xlsx,
        ".csv":  extract_text_from_csv,
    }
    
    fn = dispatch.get(extension)
    if not fn:
        return None, extension
    
    return fn(filepath), extension


# ══════════════════════════════════════════════════════════════════════════════
# ATTACHMENT HANDLING
# ══════════════════════════════════════════════════════════════════════════════

def save_attachment(filename: str, payload: bytes) -> str:
    """Saves bytes to ATTACHMENT_DIR, avoiding filename collisions."""
    os.makedirs(ATTACHMENT_DIR, exist_ok=True)
    base, ext = os.path.splitext(filename)
    dest = os.path.join(ATTACHMENT_DIR, filename)
    counter = 1
    while os.path.exists(dest):
        dest = os.path.join(ATTACHMENT_DIR, f"{base}_{counter}{ext}")
        counter += 1
    with open(dest, "wb") as f:
        f.write(payload)
    log(f"  Saved attachment: {dest} ({len(payload):,} bytes)")
    return dest


def extract_supported_attachments(msg: pyzmail.PyzMessage) -> list[str]:
    """Returns list of saved file paths for all supported document types."""
    saved_paths = []
    for part in msg.mailparts:
        if part.disposition != "attachment" and part.filename is None:
            continue
        filename = part.filename or "attachment"
        extension = os.path.splitext(filename)[1].lower()
        
        if extension not in SUPPORTED_EXTENSIONS:
            continue
            
        payload = part.get_payload()
        if not isinstance(payload, bytes):
            continue
        path = save_attachment(filename, payload)
        saved_paths.append(path)
    return saved_paths


# ══════════════════════════════════════════════════════════════════════════════
# OPENAI TABLE/COLUMN EXTRACTION + DOCUMENT BRIEF
# ══════════════════════════════════════════════════════════════════════════════

OPENAI_TABLE_SYSTEM = (
    "You are a document analysis assistant. "
    "Identify every table in a document, show each as a Markdown table, "
    "then output a single JSON block with all column headers grouped by table. "
    "No extra explanation outside the analysis and the JSON block."
)

OPENAI_TABLE_USER = """Below is the plain-text content extracted from a document.
Tables in DOCX files appear with cells separated by ' | ' and rows on separate lines.

Your task:
  1. Identify EVERY table present in the text.
  2. For each table:
     - Number it sequentially (Table 1, Table 2, …).
     - Include any nearby title or heading.
     - Show column headers.
     - Show ALL data rows as a clean Markdown table.
  3. After the human-readable analysis, output a JSON block in this EXACT format
     (no other text after it):

```json
{{
  "columns_by_table": {{
    "table_1": ["Col A", "Col B"],
    "table_2": ["Col X", "Col Y"]
  }}
}}
```

  4. Do NOT invent data. Only describe what is actually in the text.
  5. If no tables are found return: {{"columns_by_table": {{}}}}

--- DOCUMENT TEXT START ---
{text_content}
--- DOCUMENT TEXT END ---

Now provide the analysis followed by the JSON block.
"""

OPENAI_BRIEF_SYSTEM = "You are a document analysis assistant. Analyze documents and return structured JSON summaries. Return only valid JSON with no explanation, markdown, or code blocks."

OPENAI_BRIEF_USER = """This is the extracted text content from a document.

Your task:
1. Identify what TYPE OF INFORMATION this document contains (e.g., "arrest records", "parking violations", "financial transactions", "meeting minutes", "investigation reports")
2. Identify what DATA ELEMENTS are present (e.g., "accused names", "case numbers", "dates", "amounts", "vehicle numbers")
3. Extract key entities that would help search for similar documents
4. Provide search keywords that describe the CONTENT DOMAIN (not the document title)
5. Return ONLY a JSON object - no explanation, no markdown, no code blocks.

CRITICAL: Focus on WHAT KIND OF DATA the document contains, NOT what the document is asking for or announcing.

Example output for a crime statistics report:
{{
  "document_type": "Crime Statistics Report",
  "data_content_summary": "Contains arrest records with accused details, case types, investigation officers, police station assignments, and crime categories for Zone-2 stations during January 2026",
  "key_entities": {{
    "dates": ["2026-01-01", "2026-01-02"],
    "locations": ["Grant Road", "Malabar Hill", "D.B. Marg"],
    "reference_numbers": ["GR 123/2026", "FIR 18"],
    "people": ["API Pawar", "PSI Suryawanshi"]
  }},
  "main_topics": ["arrested accused", "crime types", "police stations", "investigation officers", "case numbers", "NDPS", "women crimes", "cybercrime"],
  "data_domain": ["police records", "criminal cases", "arrest information", "investigation data"]
}}

Example output for a meeting circular:
{{
  "document_type": "Meeting Circular",
  "data_content_summary": "Contains meeting schedule, attendance requirements, and data submission deadlines - does NOT contain actual accused records or case data",
  "key_entities": {{
    "dates": ["2026-02-21"],
    "locations": ["South Control Room", "Mumbai"],
    "reference_numbers": ["33/2026"],
    "people": ["Rajeshkumar Gatthe"]
  }},
  "main_topics": ["meeting announcement", "deadline", "attendance", "submission requirements"],
  "data_domain": ["administrative", "procedural", "meeting coordination"]
}}

TEXT CONTENT:
{text_content}
"""


def extract_columns_from_text(text: str) -> tuple[list[str], dict[str, list[str]]]:
    """
    Send extracted text to OpenAI to find tables and extract column headers.

    Returns
    -------
    flat_columns   : deduplicated list of all headers across all tables
    columns_by_table : {"table_1": [...], "table_2": [...], ...}  ← kept separate
    """
    try:
        max_chars = 50000
        if len(text) > max_chars:
            log(f"    Truncating text from {len(text):,} to {max_chars:,} chars")
            text = text[:max_chars] + "\n...[truncated]"

        log(f"    Sending to OpenAI to find tables...")
        response = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": OPENAI_TABLE_SYSTEM},
                {"role": "user",   "content": OPENAI_TABLE_USER.format(text_content=text)},
            ],
            temperature=0,
        )
        raw = response.choices[0].message.content.strip()

        # Extract the JSON block (```json … ``` or bare { … })
        json_str = None
        if "```json" in raw:
            start = raw.find("```json") + 7
            end   = raw.find("```", start)
            if end != -1:
                json_str = raw[start:end].strip()
        if not json_str:
            start = raw.rfind("{")
            end   = raw.rfind("}")
            if start != -1 and end != -1:
                json_str = raw[start:end + 1]

        if not json_str:
            log(f"    No JSON block found — no tables detected.")
            return [], {}

        parsed = json.loads(json_str)
        columns_by_table: dict[str, list[str]] = parsed.get("columns_by_table", {})

        if not columns_by_table:
            log(f"    No tables found.")
            return [], {}

        # Log each table separately
        flat_columns: list[str] = []
        seen: set[str] = set()
        for table_key, headers in columns_by_table.items():
            log(f"    ✓ {table_key}: {headers}")
            for h in headers:
                if h not in seen:
                    flat_columns.append(h)
                    seen.add(h)

        log(f"    → {len(columns_by_table)} table(s) | {len(flat_columns)} unique columns total")
        return flat_columns, columns_by_table

    except Exception as e:
        log(f"    OpenAI API error: {e}")
        return [], {}


def get_document_brief(text: str) -> dict | None:
    """Get a brief summary and analysis of the document from OpenAI."""
    try:
        max_chars = 50000
        if len(text) > max_chars:
            log(f"    Truncating text for brief analysis")
            text = text[:max_chars] + "\n...[truncated]"
        
        log(f"    Getting document brief from OpenAI...")
        response = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": OPENAI_BRIEF_SYSTEM},
                {"role": "user", "content": OPENAI_BRIEF_USER.format(text_content=text)},
            ],
            temperature=0,
        )
        raw = response.choices[0].message.content.strip()
        
        if raw.startswith("```"):
            raw = raw.replace("```json", "").replace("```", "").strip()
        start, end = raw.find("{"), raw.rfind("}")
        if start == -1 or end == -1:
            log(f"    Could not get document brief")
            return None
        
        brief = json.loads(raw[start:end + 1])
        log(f"    ✓ Document type: {brief.get('document_type', 'Unknown')}")
        log(f"    ✓ Data content: {brief.get('data_content_summary', 'N/A')[:100]}")
        log(f"    ✓ Main topics: {brief.get('main_topics', [])}")
        return brief
        
    except Exception as e:
        log(f"    Brief analysis error: {e}")
        return None


# ══════════════════════════════════════════════════════════════════════════════
# OPENCLAW INTEGRATION (ORIGINAL LOGIC - NO CHANGES)
# ══════════════════════════════════════════════════════════════════════════════

def call_openclaw(subject: str, sender: str, body: str,
                  columns_by_table: dict[str, list[str]],
                  document_brief: dict | None = None) -> dict | None:
    """
    Sends email context + per-table column headers + document brief to OpenClaw.
    Each table is sent as a separate entry so OpenClaw can match and return
    data for each table independently.

    FIX-3  : DB_PATH is passed dynamically so the skill never uses a stale
              hardcoded path.  The skill writes temp results to a unique
              RUN_ID-stamped file to avoid race conditions (FIX-5).
    FIX-7  : Only a 3 000-char excerpt of each brief is forwarded to keep
              the total prompt within a safe token budget.
    FIX-1  : A hard JSON-only footer is appended to every prompt.
    """

    # ── Unique run identifier stops concurrent emails clobbering each other ──
    run_id = str(uuid.uuid4()).replace("-", "")[:16]   # 16-char hex string

    brief_section = ""
    if document_brief:
        # Trim long summaries to keep token budget under control (FIX-7)
        data_summary = (document_brief.get('data_content_summary', 'N/A') or '')[:500]
        brief_section = f"""
DOCUMENT BRIEF (from OpenAI analysis):
Document Type: {document_brief.get('document_type', 'Unknown')}
Data Content: {data_summary}
Main Topics: {', '.join((document_brief.get('main_topics', []) or [])[:10])}
Data Domain: {', '.join((document_brief.get('data_domain', []) or [])[:5])}
Key Entities:
  - Dates: {', '.join((document_brief.get('key_entities', {}) or {}).get('dates', [])[:5])}
  - Locations: {', '.join((document_brief.get('key_entities', {}) or {}).get('locations', [])[:5])}
  - Reference Numbers: {', '.join((document_brief.get('key_entities', {}) or {}).get('reference_numbers', [])[:5])}
  - People: {', '.join((document_brief.get('key_entities', {}) or {}).get('people', [])[:5])}
"""

    # Format each table on its own line so OpenClaw sees them as distinct
    tables_section = "\n".join(
        f"  {tbl}: {json.dumps(cols, ensure_ascii=False)}"
        for tbl, cols in columns_by_table.items()
    )

    # FIX-1 / FIX-3 : inject DB path + run_id; hard JSON-only footer
    prompt = f"""Use the email-context-responder skill.

DATABASE PATH: {DB_PATH}
RUN_ID: {run_id}

DETECTED TABLES & COLUMNS (from attachment — each table is separate):
{tables_section}
{brief_section}
EMAIL:
Subject : {json.dumps(subject)}
From    : {json.dumps(sender)}
Body    :
{body}

⚠ RESPOND WITH ONLY VALID JSON. DO NOT WRITE ANYTHING BEFORE THE OPENING '{{' OR AFTER THE CLOSING '}}'.
"""

    payload = {
        "model": "openclaw",
        "messages": [{"role": "user", "content": prompt}]
    }
    headers = {
        "Authorization": f"Bearer {TOKEN}",
        "Content-Type": "application/json",
        "x-openclaw-agent-id": "main"
    }

    try:
        resp = requests.post(OPENCLAW_URL, headers=headers, json=payload, timeout=120)
        resp.raise_for_status()
        raw = resp.json()["choices"][0]["message"]["content"]
        log("  --- OpenClaw response (first 500 chars) ---")
        log(raw[:500])
        log("  -------------------------------------------")
        return _parse_json(raw)
    except Exception as e:
        log(f"  OpenClaw API error: {e}")
        return None


def _parse_json(raw: str) -> dict | None:
    """
    Robustly extract the outermost JSON object from an LLM response.

    FIX-1: Uses a regex-based approach that correctly handles:
      - Markdown code fences  (```json ... ```)
      - Preamble text before the '{'
      - Trailing prose after the '}'
      - Nested JSON inside the object (rfind was breaking this before)
    """
    try:
        cleaned = raw.strip()

        # Strip markdown fences first
        cleaned = re.sub(r"^```json\s*", "", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"^```\s*",    "", cleaned)
        cleaned = re.sub(r"```\s*$",    "", cleaned)
        cleaned = cleaned.strip()

        # Find the outermost {...} by counting braces (handles nested objects)
        start = cleaned.find("{")
        if start == -1:
            raise ValueError("No JSON object found in response.")

        depth   = 0
        in_str  = False
        escape  = False
        end_idx = -1

        for i, ch in enumerate(cleaned[start:], start=start):
            if escape:
                escape = False
                continue
            if ch == "\\" and in_str:
                escape = True
                continue
            if ch == '"':
                in_str = not in_str
                continue
            if in_str:
                continue
            if ch == "{":
                depth += 1
            elif ch == "}":
                depth -= 1
                if depth == 0:
                    end_idx = i
                    break

        if end_idx == -1:
            raise ValueError("Unmatched braces — incomplete JSON object.")

        return json.loads(cleaned[start:end_idx + 1])

    except (json.JSONDecodeError, ValueError) as e:
        log(f"  JSON parse error: {e}")
        log(f"  Raw response snippet: {raw[:300]}")
        return None


# ══════════════════════════════════════════════════════════════════════════════
# REPLY DOCUMENT CREATION - SAME FORMAT AS RECEIVED
# ══════════════════════════════════════════════════════════════════════════════

def create_reply_pdf(tables: list[dict], subject: str, output_path: str,
                     facts: list[str] | None = None) -> str:
    """
    Creates a styled PDF with ONE section per table.
    Each table gets its own heading + grid.
    Uses NotoSansDevanagari per-cell so Hindi text renders correctly.

    `tables` is a list of dicts, each with:
        {
            "title":   "Table 1 — Arrested Accused",   # optional
            "headers": ["Col A", "Col B", ...],
            "rows":    [["val1", "val2", ...], ...]
        }
    """
    _load_hindi_fonts()

    if not tables:
        raise ValueError("No tables provided for PDF creation.")

    # Use landscape if any table has more than 5 columns
    max_cols   = max((len(t.get("headers", [])) for t in tables), default=0)
    page_size  = landscape(A4) if max_cols > 5 else A4

    doc = SimpleDocTemplate(
        output_path, pagesize=page_size,
        leftMargin=1.5 * cm, rightMargin=1.5 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    story  = []

    # ── Document title ──────────────────────────────────────────────────────
    title_style = ParagraphStyle("CustomTitle", parent=styles["Title"],
                                 fontSize=14, spaceAfter=12)
    story.append(Paragraph(f"Response: {subject}", title_style))
    story.append(Paragraph(
        f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        styles["Normal"]
    ))
    story.append(Spacer(1, 0.5 * cm))

    # ── Shared cell styles (Latin + Hindi) ──────────────────────────────────
    cell_latin = ParagraphStyle("CellLatin", parent=styles["Normal"],
                                fontSize=8, leading=11, fontName="Helvetica")
    cell_hindi = ParagraphStyle("CellHindi", parent=styles["Normal"],
                                fontSize=8, leading=11, fontName=_hindi_font_regular)
    hdr_latin  = ParagraphStyle("HdrLatin", parent=styles["Normal"],
                                fontSize=8, leading=11,
                                textColor=colors.white, fontName="Helvetica-Bold")
    hdr_hindi  = ParagraphStyle("HdrHindi", parent=styles["Normal"],
                                fontSize=8, leading=11,
                                textColor=colors.white, fontName=_hindi_font_bold)
    tbl_title_style = ParagraphStyle("TblTitle", parent=styles["Heading2"],
                                     fontSize=11, spaceAfter=6)

    usable_width = page_size[0] - 3 * cm

    # ── One section per table ───────────────────────────────────────────────
    for idx, tbl_data in enumerate(tables, start=1):
        tbl_title = tbl_data.get("title") or f"Table {idx}"
        headers   = tbl_data.get("headers", [])
        rows      = tbl_data.get("rows",    [])

        if not headers:
            log(f"  ⚠ Table {idx} has no headers — skipping.")
            continue

        story.append(Paragraph(tbl_title, tbl_title_style))

        # Build rows
        table_body = [[_smart_paragraph(str(h), hdr_latin, hdr_hindi) for h in headers]]
        for row in rows:
            padded = list(row) + [""] * (len(headers) - len(row))
            padded = padded[:len(headers)]
            table_body.append([
                _smart_paragraph(str(cell), cell_latin, cell_hindi) for cell in padded
            ])

        col_width = usable_width / len(headers)
        tbl = Table(table_body, colWidths=[col_width] * len(headers), repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND",     (0, 0), (-1, 0), colors.HexColor("#2C3E50")),
            ("TEXTCOLOR",      (0, 0), (-1, 0), colors.white),
            ("FONTNAME",       (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F2F2")]),
            ("GRID",           (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
            ("VALIGN",         (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING",     (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING",  (0, 0), (-1, -1), 4),
            ("LEFTPADDING",    (0, 0), (-1, -1), 4),
            ("RIGHTPADDING",   (0, 0), (-1, -1), 4),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 0.6 * cm))   # gap between tables

    # ── Source facts ─────────────────────────────────────────────────────────
    if facts:
        story.append(Spacer(1, 0.4 * cm))
        story.append(Paragraph("Source Facts", styles["Heading2"]))
        story.append(Spacer(1, 0.2 * cm))
        fact_latin = ParagraphStyle("FactLatin", parent=styles["Normal"],
                                    fontSize=8, leading=12, leftIndent=10,
                                    fontName="Helvetica")
        fact_hindi = ParagraphStyle("FactHindi", parent=styles["Normal"],
                                    fontSize=8, leading=12, leftIndent=10,
                                    fontName=_hindi_font_regular)
        for i, fact in enumerate(facts, start=1):
            story.append(_smart_paragraph(f"{i}. {fact}", fact_latin, fact_hindi))

    doc.build(story)
    log(f"  Reply PDF created: {output_path} ({len(tables)} table(s))")
    return output_path


def create_reply_docx(tables: list[dict], subject: str, output_path: str,
                      facts: list[str] | None = None) -> str:
    """Creates a DOCX with one section per table."""
    if not tables:
        raise ValueError("No tables provided for DOCX creation.")

    doc = docx.Document()

    title = doc.add_heading(f"Response: {subject}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    for idx, tbl_data in enumerate(tables, start=1):
        tbl_title = tbl_data.get("title") or f"Table {idx}"
        headers   = tbl_data.get("headers", [])
        rows      = tbl_data.get("rows",    [])

        if not headers:
            continue

        doc.add_paragraph()
        doc.add_heading(tbl_title, level=2)

        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = str(header)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)

        for row_idx, row_data in enumerate(rows, start=1):
            row_cells = table.rows[row_idx].cells
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(headers):
                    row_cells[col_idx].text = str(cell_data)

    if facts:
        doc.add_paragraph()
        doc.add_heading("Source Facts", level=2)
        for i, fact in enumerate(facts, start=1):
            doc.add_paragraph(f"{i}. {fact}", style='List Number')

    doc.save(output_path)
    log(f"  Reply DOCX created: {output_path} ({len(tables)} table(s))")
    return output_path


def create_reply_xlsx(tables: list[dict], subject: str, output_path: str,
                      facts: list[str] | None = None) -> str:
    """Creates an XLSX with one sheet per table."""
    if not tables:
        raise ValueError("No tables provided for XLSX creation.")

    wb = openpyxl.Workbook()
    wb.remove(wb.active)   # remove default empty sheet

    for idx, tbl_data in enumerate(tables, start=1):
        tbl_title = tbl_data.get("title") or f"Table {idx}"
        headers   = tbl_data.get("headers", [])
        rows      = tbl_data.get("rows",    [])

        if not headers:
            continue

        # Sheet name max 31 chars, strip illegal chars
        sheet_name = tbl_title[:31].replace("/", "-").replace("\\", "-").replace("*", "").replace("?", "").replace("[", "").replace("]", "").replace(":", "")
        ws = wb.create_sheet(title=sheet_name)

        ws.append([f"Response: {subject}"])
        ws.append([f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"])
        ws.append([tbl_title])
        ws.append([])

        header_row_num = ws.max_row + 1
        ws.append(headers)
        for cell in ws[header_row_num]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")

        for row_data in rows:
            padded = list(row_data) + [""] * (len(headers) - len(row_data))
            ws.append(padded[:len(headers)])

        for column in ws.columns:
            max_length = 0
            col_letter = column[0].column_letter
            for cell in column:
                try:
                    if cell.value and len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            ws.column_dimensions[col_letter].width = min(max_length + 2, 50)

    if facts:
        ws_facts = wb.create_sheet(title="Source Facts")
        ws_facts.append(["Source Facts"])
        ws_facts.append([])
        for i, fact in enumerate(facts, start=1):
            ws_facts.append([f"{i}. {fact}"])

    wb.save(output_path)
    log(f"  Reply XLSX created: {output_path} ({len(tables)} sheet(s))")
    return output_path


def create_reply_csv(tables: list[dict], output_path: str) -> str:
    """Creates a CSV with all tables stacked, separated by a blank row."""
    if not tables:
        raise ValueError("No tables provided for CSV creation.")

    with open(output_path, "w", newline="", encoding="utf-8") as f:
        writer = csv_module.writer(f)
        for idx, tbl_data in enumerate(tables, start=1):
            tbl_title = tbl_data.get("title") or f"Table {idx}"
            headers   = tbl_data.get("headers", [])
            rows      = tbl_data.get("rows",    [])

            if not headers:
                continue

            if idx > 1:
                writer.writerow([])   # blank separator between tables

            writer.writerow([tbl_title])
            writer.writerow(headers)
            for row_data in rows:
                padded = list(row_data) + [""] * (len(headers) - len(row_data))
                writer.writerow(padded[:len(headers)])

    log(f"  Reply CSV created: {output_path} ({len(tables)} table(s))")
    return output_path


def create_reply_document(tables: list[dict], subject: str, original_extension: str,
                          facts: list[str] | None = None) -> str:
    """
    Creates reply document in the SAME FORMAT as the original attachment.
    PDF → PDF, DOCX → DOCX, XLSX → XLSX, CSV → CSV
    All formats now support multiple tables.
    """
    os.makedirs(ATTACHMENT_DIR, exist_ok=True)
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

    dispatch = {
        ".pdf":  (create_reply_pdf,  f"reply_{ts}.pdf"),
        ".docx": (create_reply_docx, f"reply_{ts}.docx"),
        ".xlsx": (create_reply_xlsx, f"reply_{ts}.xlsx"),
        ".csv":  (create_reply_csv,  f"reply_{ts}.csv"),
    }

    creator_fn, filename = dispatch.get(original_extension, (create_reply_pdf, f"reply_{ts}.pdf"))
    output_path = os.path.join(ATTACHMENT_DIR, filename)

    if original_extension == ".csv":
        return creator_fn(tables, output_path)
    else:
        return creator_fn(tables, subject, output_path, facts)


# ══════════════════════════════════════════════════════════════════════════════
# EMAIL REPLY
# ══════════════════════════════════════════════════════════════════════════════

def send_reply_with_attachment(to_email: str, subject: str,
                                body_text: str, attachment_path: str) -> None:
    """Send reply email with attachment."""
    try:
        msg = MIMEMultipart()
        msg["From"]    = SMTP_EMAIL
        msg["To"]      = to_email
        msg["Subject"] = f"Re: {subject}"

        msg.attach(MIMEText(body_text, "plain", "utf-8"))

        with open(attachment_path, "rb") as f:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(f.read())
        encoders.encode_base64(part)
        attachment_name = os.path.basename(attachment_path)
        part.add_header("Content-Disposition", f'attachment; filename="{attachment_name}"')
        msg.attach(part)

        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(SMTP_EMAIL, SMTP_PASSWORD)
            server.sendmail(SMTP_EMAIL, to_email, msg.as_string())

        log(f"  Reply sent to {to_email} with attachment: {attachment_name}")
    except Exception as e:
        log(f"  SMTP error: {e}")


def send_text_only_reply(to_email: str, subject: str, body: str) -> None:
    """Send text-only reply email."""
    try:
        msg = MIMEMultipart()
        msg["From"]    = SMTP_EMAIL
        msg["To"]      = to_email
        msg["Subject"] = f"Re: {subject}"
        msg.attach(MIMEText(body, "plain", "utf-8"))
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(SMTP_EMAIL, SMTP_PASSWORD)
            server.sendmail(SMTP_EMAIL, to_email, msg.as_string())
        log(f"  Text-only reply sent to {to_email}")
    except Exception as e:
        log(f"  SMTP error: {e}")


# ══════════════════════════════════════════════════════════════════════════════
# MAIN EMAIL HANDLER
# ══════════════════════════════════════════════════════════════════════════════

def on_new_email(subject: str, sender: str, body: str,
                 attachment_paths: list[str]) -> None:
    """Process new email with attachments."""
    log("="*70)
    log("NEW EMAIL")
    log("="*70)
    log(f"From    : {sender}")
    log(f"Subject : {subject}")
    log(f"Preview : {body[:120].replace(chr(10), ' ')}")
    log(f"Attachments ({len(attachment_paths)}): {[os.path.basename(p) for p in attachment_paths]}")
    log("="*70)

    processed = load_processed()
    all_columns: list[str] = []
    all_columns_by_table: dict[str, list[str]] = {}   # ← keeps tables separate
    all_briefs = []
    original_extension = None
    
    for file_path in attachment_paths:
        filename = os.path.basename(file_path)
        
        if filename in processed:
            log(f"  Already processed '{filename}' — skipping.")
            continue
        
        log(f"\n  Processing: {filename}")
        log(f"  {'─'*66}")
        
        extracted_text, extension = extract_document_content(file_path)
        
        if original_extension is None:
            original_extension = extension
        
        if extracted_text is None:
            log(f"    ✗ Extraction failed")
            continue
        
        if not extracted_text.strip():
            log(f"    ⚠ No text extracted (empty document)")
            continue
        
        log(f"    ✓ Extracted {len(extracted_text):,} characters")
        
        preview = extracted_text[:150].replace('\n', ' ')[:120]
        log(f"    Preview: {preview}...")
        
        brief = get_document_brief(extracted_text)
        if brief:
            all_briefs.append(brief)
        
        flat_cols, cols_by_table = extract_columns_from_text(extracted_text)

        # Merge into the running dicts, namespacing keys by filename to avoid
        # collisions when multiple attachments are present
        # e.g.  "report.pdf_table_1", "report.pdf_table_2"
        for tbl_key, headers in cols_by_table.items():
            namespaced_key = f"{filename}_{tbl_key}"
            all_columns_by_table[namespaced_key] = headers

        all_columns.extend(flat_cols)
        
        mark_processed(filename)
    
    if not all_columns_by_table:
        log("\n  No tables/columns found in any attachment — skipping.")
        return
    
    # Unique flat list (for logging only)
    seen: set[str] = set()
    unique_columns: list[str] = []
    for col in all_columns:
        if col not in seen:
            unique_columns.append(col)
            seen.add(col)
    
    log(f"\n  Total tables found   : {len(all_columns_by_table)}")
    log(f"  Total unique columns : {len(unique_columns)}")
    
    combined_brief = None
    if all_briefs:
        all_topics = []
        all_dates = []
        all_locations = []
        all_ref_nums = []
        all_people = []
        summaries = []
        
        for brief in all_briefs:
            summaries.append(brief.get('brief_summary', ''))
            all_topics.extend(brief.get('main_topics', []))
            entities = brief.get('key_entities', {})
            all_dates.extend(entities.get('dates', []))
            all_locations.extend(entities.get('locations', []))
            all_ref_nums.extend(entities.get('reference_numbers', []))
            all_people.extend(entities.get('people', []))
        
        combined_brief = {
            'document_type': all_briefs[0].get('document_type', 'Unknown'),
            'brief_summary': ' '.join(summaries),
            'main_topics': list(set(all_topics)),
            'key_entities': {
                'dates': list(set(all_dates)),
                'locations': list(set(all_locations)),
                'reference_numbers': list(set(all_ref_nums)),
                'people': list(set(all_people))
            }
        }
        
        log(f"\n  Document Brief:")
        log(f"    Type: {combined_brief['document_type']}")
        log(f"    Topics: {combined_brief['main_topics']}")
    
    log("\n  Calling OpenClaw...")
    result = call_openclaw(subject, sender, body, all_columns_by_table, combined_brief)
    
    if not result:
        log("  OpenClaw analysis failed — skipping.")
        return
    
    log(f"\n  Category : {result.get('category', 'unknown')}")
    log(f"  Priority : {result.get('priority', 'unknown')}")
    log(f"  Language : {result.get('language_of_reply', 'unknown')}")
    
    matched = result.get("matched_documents", [])
    if matched:
        log(f"  Matched docs ({len(matched)}): {[d.get('file_name') for d in matched]}")
    
    if result.get("reply_note"):
        log(f"  Note     : {result['reply_note']}")
    
    requires_reply  = result.get("requires_reply", False)
    suggested_reply = result.get("suggested_reply", "").strip()
    facts           = result.get("data_used_in_reply", {}).get("facts", [])

    if not requires_reply:
        log("\n  No reply required per OpenClaw.")
        return

    # ── Collect all tables from OpenClaw response ───────────────────────────
    # OpenClaw may return:
    #   "tables": [{"title":..., "headers":..., "rows":...}, ...]   ← preferred
    #   "table_data": {"headers":..., "rows":...}                   ← legacy single
    tables: list[dict] = result.get("tables", [])

    if not tables:
        # Legacy: single table_data → wrap in a list
        td = result.get("table_data", {})
        if td and td.get("headers"):
            tables = [td]

    log(f"\n  Tables in OpenClaw response: {len(tables)}")
    for i, t in enumerate(tables, 1):
        log(f"    Table {i}: '{t.get('title', '(no title)')}' — {len(t.get('headers', []))} cols, {len(t.get('rows', []))} rows")

    if not tables:
        log("\n  No table data — sending text-only reply.")
        if suggested_reply:
            send_text_only_reply(sender, subject, suggested_reply)
        return

    log(f"\n  Creating reply in format: {original_extension}")

    try:
        reply_path = create_reply_document(
            tables,
            subject,
            original_extension or ".pdf",
            facts=facts
        )
    except Exception as e:
        log(f"\n  Document creation failed: {e} — sending text-only reply.")
        if suggested_reply:
            send_text_only_reply(sender, subject, suggested_reply)
        return

    reply_body = suggested_reply or (
        "Please find the extracted data attached.\n\nRegards,\nAuto-Reply System"
    )
    send_reply_with_attachment(sender, subject, reply_body, reply_path)
    log("\n" + "="*70)


# ══════════════════════════════════════════════════════════════════════════════
# IMAP INBOX POLLING
# ══════════════════════════════════════════════════════════════════════════════

def check_for_new_mail() -> None:
    """Check for new unread emails and process them."""
    with imaplib.IMAP4_SSL(IMAP_HOST) as server:
        server.login(EMAIL, PASSWORD)
        server.select("INBOX")

        today = datetime.date.today().strftime("%d-%b-%Y")
        _, uids = server.search(None, f'(UNSEEN SINCE "{today}")')

        if not uids or not uids[0]:
            log("No new emails.")
            return

        uid_list = uids[0].split()
        log(f"Found {len(uid_list)} new email(s).")

        for uid in uid_list:
            _, raw_data = server.fetch(uid, "(RFC822)")
            if not raw_data or not raw_data[0]:
                continue

            msg = pyzmail.PyzMessage.factory(raw_data[0][1])

            subject   = msg.get_subject() or "(No Subject)"
            addresses = msg.get_addresses("from")
            sender    = addresses[0][1] if addresses else "unknown"

            if msg.text_part:
                charset = msg.text_part.charset or "utf-8"
                body = msg.text_part.get_payload().decode(charset, errors="ignore")
            else:
                body = "(No text body)"

            attachment_paths = extract_supported_attachments(msg)

            if attachment_paths:
                on_new_email(subject, sender, body, attachment_paths)
            else:
                log(f"No supported attachments in email from {sender} — skipping.")

            server.store(uid, "+FLAGS", "\\Seen")


# ══════════════════════════════════════════════════════════════════════════════
# MAIN LOOP
# ══════════════════════════════════════════════════════════════════════════════

def main() -> None:
    """Main polling loop."""
    log("="*70)
    log("ENHANCED EMAIL PIPELINE - Same Format Reply")
    log("="*70)
    log(f"Supported types : {', '.join(SUPPORTED_EXTENSIONS)}")
    log(f"Attachment dir  : {ATTACHMENT_DIR}/")
    log(f"Poll interval   : {CHECK_INTERVAL}s")
    log(f"Krutidev support: {'✓ Available' if KRUTIDEV_AVAILABLE else '✗ Not available'}")
    log(f"AI provider     : OpenAI ({OPENAI_MODEL})")
    log(f"Reply format    : SAME as received (PDF→PDF, DOCX→DOCX, etc.)")

    # Pre-load Hindi font at startup so any warnings appear immediately
    _load_hindi_fonts()
    log(f"Hindi PDF font  : {_hindi_font_regular}")
    log("="*70)
    log("")

    while True:
        try:
            check_for_new_mail()
        except Exception as e:
            log(f"Error in main loop: {e}")
        time.sleep(CHECK_INTERVAL)


if __name__ == "__main__":
    main()